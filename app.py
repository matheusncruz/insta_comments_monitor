import os
import requests
import re
import time
import streamlit as st
import pandas as pd
from datetime import datetime
from collections import defaultdict
from docx import Document

# **📌 Interface Principal**
st.title("📊 Monitoramento de Engajamento no Instagram")
st.write("Escolha o tipo de análise que deseja realizar.")

# **Seleção do Tipo de Análise**
analysis_type = st.radio("📌 Escolha uma opção:", ["Monitoramento de Palavras", "Indicadores por Mês"])

# **Inputs do Usuário**
ACCESS_TOKEN = st.text_input("🔑 Token de Acesso (Graph API)", type="password")
IG_ACCOUNT_ID = st.text_input("🆔 ID da Conta do Instagram")

# **📌 Criando lista fixa de meses**
def generate_fixed_months():
    fixed_months = []
    for year in range(2022, 2028):
        for month in range(1, 13):
            fixed_months.append(f"{month:02d}/{year}")
    return fixed_months

available_months = generate_fixed_months()

# **Escolha do Tipo de Pesquisa**
if analysis_type == "Monitoramento de Palavras":
    filter_type = st.radio("📊 Como deseja buscar os posts?", ["Últimos X posts", "Selecionar meses específicos"])
    
    if filter_type == "Últimos X posts":
        num_posts = st.number_input("📌 Quantidade de Posts a Analisar", min_value=1, max_value=500)
        selected_months = []
    else:
        num_posts = None
        selected_months = st.multiselect("📅 Selecione os meses (formato MM/YYYY)", available_months)

    keywords = st.text_area("🔍 Palavras-Chave (separe por vírgula)", "")

elif analysis_type == "Indicadores por Mês":
    filter_type = st.radio("📊 Como deseja buscar os posts?", ["Últimos X posts", "Selecionar meses específicos"])
    
    if filter_type == "Últimos X posts":
        num_posts = st.number_input("📌 Quantidade de Posts a Analisar", min_value=1, max_value=500)
    else:
        num_posts = None
        selected_months = st.multiselect("📅 Selecione os meses (formato MM/YYYY)", available_months)  

# **📌 Botão para rodar o script**
if st.button("🚀 Executar Análise"):
    if not ACCESS_TOKEN or not IG_ACCOUNT_ID:
        st.error("⚠ Preencha o Token de Acesso e o ID da Conta do Instagram.")
    else:
        def log(text):
            """ Exibe logs na interface para depuração """
            st.text(text)

        # **📌 Coletar Posts**
        post_list = []
        monthly_data = defaultdict(lambda: {"likes": 0, "comments": 0, "impressions": 0, "reach": 0})

        url_posts = f"https://graph.facebook.com/v18.0/{IG_ACCOUNT_ID}/media?fields=id,permalink,timestamp,like_count,comments_count&limit=100&access_token={ACCESS_TOKEN}"
        
        while (num_posts and len(post_list) < num_posts) or (not num_posts and url_posts):
            response = requests.get(url_posts)
            if response.status_code == 200:
                posts_data = response.json()
                for post in posts_data["data"]:
                    date_obj = datetime.strptime(post["timestamp"], "%Y-%m-%dT%H:%M:%S%z")
                    post_month = date_obj.strftime("%m/%Y")
                    
                    if filter_type == "Selecionar meses específicos" and post_month not in selected_months:
                        continue
                    
                    post_data = {
                        "id": post["id"],
                        "link": post["permalink"],
                        "date": post_month,
                        "likes": post.get("like_count", 0),
                        "comments": post.get("comments_count", 0),
                        "impressions": 0,
                        "reach": 0
                    }

                    # **📌 Buscar Impressões e Alcance do Post**
                    url_insights = f"https://graph.facebook.com/v18.0/{post['id']}/insights?metric=impressions,reach&access_token={ACCESS_TOKEN}"
                    response_insights = requests.get(url_insights)
                    if response_insights.status_code == 200:
                        insights_data = response_insights.json()
                        for metric in insights_data.get("data", []):
                            if metric["name"] == "impressions":
                                post_data["impressions"] = int(metric["values"][0]["value"])
                            elif metric["name"] == "reach":
                                post_data["reach"] = int(metric["values"][0]["value"])

                    post_list.append(post_data)
                    monthly_data[post_month]["likes"] += post_data["likes"]
                    monthly_data[post_month]["comments"] += post_data["comments"]
                    monthly_data[post_month]["impressions"] += post_data["impressions"]
                    monthly_data[post_month]["reach"] += post_data["reach"]

                    if num_posts and len(post_list) >= num_posts:
                        break
                
                url_posts = posts_data.get("paging", {}).get("next")

        log(f"\n✅ Total de {len(post_list)} posts coletados!\n")

        # **📌 Criar DataFrame e Mostrar Indicadores**
        if analysis_type == "Indicadores por Mês":
            df = pd.DataFrame.from_dict(monthly_data, orient="index").reset_index().rename(columns={"index": "Mês"})
            st.subheader("📊 Indicadores Mensais")
            st.write(df)

            doc = Document()
            doc.add_heading("Relatório - Indicadores Mensais", level=1)

            for index, row in df.iterrows():
                doc.add_paragraph(
                    f"📅 {row['Mês']}\n"
                    f"👍 {row['likes']} curtidas | 💬 {row['comments']} comentários | 👀 {row['impressions']} visualizações | 🎯 {row['reach']} alcance"
                )

            word_filename = "relatorio_indicadores.docx"
            doc.save(word_filename)
            with open(word_filename, "rb") as file:
                st.download_button("📥 Baixar Relatório Word", file, file_name=word_filename)

        # **📌 Monitoramento de Palavras**
        elif analysis_type == "Monitoramento de Palavras":
            keywords = {word.strip().lower() for word in keywords.split(",")}
            keyword_count = {word: 0 for word in keywords}
            keyword_likes = {word: 0 for word in keywords}

            doc = Document()
            doc.add_heading("Relatório - Monitoramento de Palavras", level=1)

            for post in post_list:
                url_comments = f"https://graph.facebook.com/v18.0/{post['id']}/comments?fields=id,text,username,like_count&access_token={ACCESS_TOKEN}"
                response_comments = requests.get(url_comments)

                log(f"\n📌 Analisando Post: {post['link']}")
                doc.add_heading(f"📌 Post: {post['link']}", level=2)

                if response_comments.status_code == 200:
                    comments_data = response_comments.json()
                    for comment in comments_data.get("data", []):
                        text = comment.get("text", "").lower()
                        username = comment.get("username", "Usuário desconhecido")
                        like_count = comment.get("like_count", 0)

                        for word in keywords:
                            if re.search(fr"\b{re.escape(word)}\b", text, re.IGNORECASE):
                                keyword_count[word] += 1
                                keyword_likes[word] += like_count
                                log(f"   💬 {username} ({like_count} curtidas): {text}")
                                doc.add_paragraph(f"💬 {username} ({like_count} curtidas): {text}")

            st.subheader("📊 Resultado - Monitoramento de Palavras")
            for word, count in keyword_count.items():
                st.write(f"**{word}**: {count} vezes. Total de curtidas: {keyword_likes[word]}")

            word_filename = "relatorio_monitoramento_palavras.docx"
            doc.save(word_filename)
            with open(word_filename, "rb") as file:
                st.download_button("📥 Baixar Relatório Word", file, file_name=word_filename)
